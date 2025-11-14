# streamlit_app.py
import streamlit as st
from pipeline import pipeline_complet
from display import st_afficher_table
from io_utils import lire_tous_les_cvs
import pandas as pd
import os

st.set_page_config(page_title="CV Scorer", layout="wide")

st.title("📋 CV Scorer — Compare CVs vs Offre")

with st.sidebar:
    st.header("Paramètres")
    folder = st.text_input("Dossier contenant les CV (.docx)", value="./CVs")
    st.caption("Indique un chemin local (ex: ./CVs). Si vide ou invalide, utiliser l'uploader ci-dessous.")
    use_uploader = st.checkbox("Uploader des CV manuellement (fallback)", value=False)
    if use_uploader:
        uploaded = st.file_uploader("Upload .docx files", accept_multiple_files=True, type=["docx"])
    st.markdown("---")
    st.write("Options de scoring")
    method = st.selectbox("Méthode de scoring", ["frequency (default)", "presence", "tfidf (optional)"])
    max_occ = st.slider("max_occurrences (pour frequency)", 1, 5, 2)

st.header("Offre de poste")
offre_text = st.text_area("Collez l'offre de poste ici (ou chargez un fichier via l'uploader)", height=220)

col1, col2 = st.columns([1,3])
with col1:
    run = st.button("▶ Lancer le scoring des CV")
with col2:
    st.write("Instructions : Indique le dossier local contenant les .docx ou coche 'Uploader' pour envoyer des fichiers. Puis coller l'offre et lancer le scoring.")

# Préparer liste original_files vide par défaut (optionnel)
original_files_input = st.text_input("Liste des fichiers originaux (séparés par virgule) — facultatif", value="")
original_files = [s.strip() for s in original_files_input.split(",") if s.strip()]

# Fallback: si uploader utilisé, sauvegarder temporairement en mémoire dict{filename:content}
if use_uploader and uploaded:
    cvs_dict = {}
    for f in uploaded:
        try:
            # streamlit UploadedFile -> bytes, passer à python-docx nécessite un fichier, mais on peut lire via Document(f)
            from docx import Document
            doc = Document(f)
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            text = ""
        cvs_dict[f.name] = text
else:
    # lire dossier
    if os.path.isdir(folder):
        cvs_dict = None  # pipeline will read folder itself
    else:
        cvs_files = st.text("Le dossier indiqué n'existe pas. Vérifiez le chemin.", value="")
        cvs_dict = {}

if run:
    if not offre_text.strip():
        st.error("Saisis l'offre de poste avant de lancer le scoring.")
    else:
        st.info("Scoring en cours... (patienter)")

        # Choix méthode
        from scoring import score_cv_frequence, score_cv_offre
        scoring_fn = score_cv_frequence
        if method == "presence":
            scoring_fn = score_cv_offre

        # Si cvs_dict non None => on a upload; sinon pipeline lit le dossier
        if cvs_dict is not None:
            # construction scores local
            scores = {}
            for fname, text in cvs_dict.items():
                if method == "frequency":
                    s = score_cv_frequence(text, offre_text, max_occurrences=max_occ)
                else:
                    s = scoring_fn(text, offre_text)
                scores[fname] = float(s)
            df = pd.DataFrame([{"Fichier": f, "Score": s} for f, s in scores.items()])
            df = df.sort_values("Score", ascending=False).reset_index(drop=True)
            # minimal grouping columns for compatibility with display
            df["base_name"] = df["Fichier"]
            df["original_score"] = None
            df["gain_vs_original"] = 0.0
        else:
            # pipeline reads folder
            df = pipeline_complet(folder=folder, original_files=original_files, offre_text=offre_text, scoring_fn=scoring_fn)
            print("Resultat de l'analyse :", df)
            
        st.success("Terminé")
        st_afficher_table(df)
